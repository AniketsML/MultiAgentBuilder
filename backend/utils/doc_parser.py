"""Document parser — extracts text from .docx, .pdf, and .txt files."""

import io
import os


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract plain text from an uploaded file based on its extension."""
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".txt":
        return file_bytes.decode("utf-8", errors="replace")

    elif ext == ".docx":
        return _extract_docx(file_bytes)

    elif ext == ".pdf":
        return _extract_pdf(file_bytes)

    elif ext in (".doc",):
        raise ValueError(
            "Legacy .doc format is not supported. "
            "Please save as .docx or export from Google Docs as .docx."
        )

    else:
        # Try to read as plain text
        try:
            return file_bytes.decode("utf-8", errors="replace")
        except Exception:
            raise ValueError(f"Unsupported file type: {ext}")


def _extract_docx(file_bytes: bytes) -> str:
    """Extract text from a .docx file."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n\n".join(paragraphs)


def _extract_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file."""
    import pdfplumber

    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

    return "\n\n".join(text_parts)
